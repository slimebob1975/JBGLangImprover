from docx import Document
from docx.shared import RGBColor
import os, sys
import json
import fitz # PyMuPDF

DEBUG = False

class JBGDocumentEditor:
    
    def __init__(self, filepath, changes):
        """
        :param filepath: Path to the input document (.docx or .pdf)
        :param changes: List of tuples like (old_text, new_text)
        """
        self.filepath = filepath
        self.changes = changes
        self.ext = os.path.splitext(filepath)[1].lower()
        self.edited_document = None

    def apply_changes(self):
        if self.ext == ".docx":
            self.edited_document = self._edit_docx()
        elif self.ext == ".pdf":
            self.edited_document = self._annotate_pdf()
        else:
            raise ValueError("Unsupported file type. Use .docx or .pdf")

    def save_edited_document(self, output_path=None):
            
        if self.edited_document:
            if self.ext == ".pdf":
                if output_path is None:
                    output_path = self.filepath.replace(".pdf", "_annotated.pdf")
                self.edited_document.save(output_path, garbage=4, deflate=True)
                self.edited_document.close()
            elif self.ext == ".docx":
                if output_path is None:
                    output_path = self.filepath.replace(".docx", "_edited.docx")
                self.edited_document.save(output_path)
            else:
                raise ValueError("Unsupported file type. Use.docx or.pdf")
            return output_path

    def _edit_docx(self):
        doc = Document(self.filepath)

        # First, handle empty paragraph insertions
        self._apply_to_empty_paragraphs(doc)
        
        for idx, para in enumerate(doc.paragraphs):
            para_index = idx + 1
            new_runs = []
            for run in para.runs:
                text = run.text
                replaced = False
                for change in self.changes:
                    if "paragraph" not in change or para_index != change["paragraph"]:
                        continue
                    old, new = change["old"], change["new"]
                    if old in text:
                        replaced = True
                        parts = text.split(old)
                        for j, part in enumerate(parts):
                            if part:
                                new_run = para.add_run(part)
                                new_runs.append(new_run)
                            if j < len(parts) - 1:
                                strike_run = para.add_run(old)
                                strike_run.font.strike = True
                                strike_run.font.color.rgb = RGBColor(255, 0, 0)
                                new_runs.append(strike_run)
                                new_run = para.add_run(new)
                                new_run.font.color.rgb = RGBColor(0, 128, 0)
                                new_runs.append(new_run)
                        break
                if not replaced:
                    new_runs.append(run)

            if new_runs:
                for _ in range(len(para.runs)):
                    para._element.remove(para.runs[0]._element)
                for run in new_runs:
                    para._element.append(run._element)

        self._suggest_nearby_paragraphs(doc)
        return doc

    def _suggest_nearby_paragraphs(self, doc):
        total = len(doc.paragraphs)
        for change in self.changes:
            if "paragraph" in change:
                para_num = change["paragraph"]
                if not (1 <= para_num <= total):
                    print(f"Warning: Paragraph {para_num} is out of range.")
                    continue
                expected_para = doc.paragraphs[para_num - 1].text
                if change["old"] not in expected_para:
                    for offset in [-2, -1, 1, 2]:
                        alt_idx = para_num - 1 + offset
                        if 0 <= alt_idx < total:
                            if change["old"] in doc.paragraphs[alt_idx].text:
                                print(f"Could not find '{change['old']}' in paragraph {para_num} — did you mean paragraph {alt_idx + 1}?")
                                break

    def _apply_to_empty_paragraphs(self, doc):
        for idx, para in enumerate(doc.paragraphs):
            para_index = idx + 1
            if para.text.strip() != "":
                continue

            for change in self.changes:
                if (
                    change.get("paragraph") == para_index
                    and change.get("old", "") == ""
                    and change.get("new")
                ):
                    para.text = change["new"]
                    print(f"Filled empty paragraph {para_index} with: {change['new']}")

    
    def _annotate_pdf_simple(self):
    
        doc = fitz.open(self.filepath)

        for change in self.changes:
            if "page" not in change or "old" not in change:
                continue

            page_index = change["page"] - 1
            if page_index >= len(doc):
                continue

            page = doc[page_index]
            rects = page.search_for(change["old"])
            if not rects:
                print(f"No match found for '{change['old']}' on page {change['page']}")
                continue

            for rect in rects:
                highlight = page.add_highlight_annot(rect)
                if "new" in change:
                    highlight.set_info(content=f"Suggestion: replace with '{change['new']}'")

        return doc
    
    def _annotate_pdf(self):
        doc = fitz.open(self.filepath)

        for change in self.changes:
            if "page" not in change or "old" not in change:
                continue

            page_index = change["page"] - 1
            if page_index >= len(doc):
                continue

            old = change["old"]
            new = change.get("new", "")
            target_line = change.get("line")
            page = doc[page_index]

            # Get all text blocks on page, sorted top-down
            blocks = sorted(page.get_text("blocks"), key=lambda b: -b[1])  # y1 descending
            line_texts = [(i + 1, b, b[4].strip()) for i, b in enumerate(blocks) if b[4].strip()]

            match_found = False

            # Try target line first
            if target_line:
                for line_no, block, text in line_texts:
                    if line_no == target_line and old in text:
                        rects = page.search_for(old, clip=fitz.Rect(block[:4]))
                        for rect in rects:
                            highlight = page.add_highlight_annot(rect)
                            if new:
                                highlight.set_info(content=f"{new}")
                        match_found = True
                        break

                # If not found, try nearby lines
                if not match_found:
                    for line_no, block, text in line_texts:
                        if abs(line_no - target_line) <= 2 and old in text:
                            rects = page.search_for(old, clip=fitz.Rect(block[:4]))
                            for rect in rects:
                                highlight = page.add_highlight_annot(rect)
                                if new:
                                    highlight.set_info(content=f"{new}")
                            print(f"Could not find '{old}' on page {change['page']}, line {target_line} — did you mean line {line_no}?")
                            match_found = True
                            break
            else:
                # Fallback if no line provided
                rects = page.search_for(old)
                for rect in rects:
                    highlight = page.add_highlight_annot(rect)
                    if new:
                        highlight.set_info(content=f"{new}")
                match_found = bool(rects)

            if not match_found:
                print(f"No match found for '{old}' on page {change['page']}.")

        return doc

def main():
    if len(sys.argv) != 3:
        print(f"Usage: python {os.path.basename(__file__)} <document_path> <changes_json_path>")
        sys.exit(1)

    doc_path = sys.argv[1]
    changes_path = sys.argv[2]

    try:
        with open(changes_path, 'r', encoding='utf-8') as f:
            changes = json.load(f)

        editor = JBGDocumentEditor(doc_path, changes)
        editor.apply_changes()
        output_path = editor.save_edited_document()
        print(f"Document processed and saved to: {output_path}")

    except Exception as e:
        print(f"An error occurred: {e}")

if __name__ == "__main__":
    main()
