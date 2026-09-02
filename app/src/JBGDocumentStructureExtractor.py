import os
import json
import docx
import sys
import posixpath
import re
from dataclasses import dataclass, asdict
from typing import Optional, Any
from lxml import etree
from zipfile import ZipFile


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
PKG_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
NSMAP = {"w": W_NS, "r": R_NS}


# ============================================================================
# Datamodeller
# ============================================================================

@dataclass
class ExtractedElement:
    type: str
    element_id: str
    text: str
    empty: bool

    part_name: str
    container_path: str

    footnote_id: Optional[str] = None
    section_index: Optional[int] = None
    header_index: Optional[int] = None
    footer_index: Optional[int] = None

    table_index: Optional[int] = None
    row_index: Optional[int] = None
    col_index: Optional[int] = None

    textbox_index: Optional[int] = None
    paragraph_index: Optional[int] = None

    story_variant: Optional[str] = None
    section_indices: Optional[list[int]] = None

    contains_linebreaks: bool = False
    contains_tabs: bool = False
    may_contain_special_runs: bool = False


# ============================================================================
# Extractor
# ============================================================================

class DocumentStructureExtractor:
    def __init__(self, filepath, logger):
        self.filepath = filepath
        self.logger = logger
        self.ext = os.path.splitext(filepath)[1].lower()
        self.structure = None

    def extract(self):
        if self.ext == ".docx":
            self.structure = self._extract_docx_structure()
        else:
            raise ValueError("Unsupported file type. Use .docx only!")
        return self.structure

    # ------------------------------------------------------------------
    # DOCX
    # ------------------------------------------------------------------

    def _extract_docx_structure(self):
        doc = docx.Document(self.filepath)
        structure = {
            "type": "docx",
            "elements": []
        }

        elements: list[ExtractedElement] = []

        # 1. Paragraphs in main document
        for i, para in enumerate(doc.paragraphs, start=1):
            text = para.text or ""
            elements.append(ExtractedElement(
                type="paragraph",
                element_id=f"paragraph_{i}",
                text=text,
                empty=not bool(text.strip()),
                part_name="word/document.xml",
                container_path=f"/document/body/paragraph[{i}]",
                paragraph_index=i,
                contains_linebreaks="\n" in text,
                contains_tabs="\t" in text,
                may_contain_special_runs=self._paragraph_may_contain_special_runs(para),
            ))

        # 2. Tables in main document
        for ti, table in enumerate(doc.tables, start=1):
            for ri, row in enumerate(table.rows, start=1):
                for ci, cell in enumerate(row.cells, start=1):
                    # A cell may contain several paragraphs.  Expose each
                    # paragraph separately so the model can never return an
                    # anchor that crosses an OOXML paragraph boundary.
                    for pi, para in enumerate(cell.paragraphs, start=1):
                        paragraph_text = para.text or ""
                        elements.append(ExtractedElement(
                            type="table_cell",
                            element_id=f"table_{ti}_cell_{ri}_{ci}_p{pi}",
                            text=paragraph_text,
                            empty=not bool(paragraph_text.strip()),
                            part_name="word/document.xml",
                            container_path=(
                                f"/document/body/table[{ti}]/row[{ri}]/"
                                f"cell[{ci}]/paragraph[{pi}]"
                            ),
                            table_index=ti,
                            row_index=ri,
                            col_index=ci,
                            paragraph_index=pi,
                            contains_linebreaks="\n" in paragraph_text,
                            contains_tabs="\t" in paragraph_text,
                            may_contain_special_runs=self._paragraph_may_contain_special_runs(para),
                        ))

        # 3. Headers and footers. Resolve the actual relationship target and
        # extract each physical story part only once, even when it is shared by
        # several sections.
        elements.extend(self._extract_header_footer_elements())

        # 4. Textboxes (main document)
        textbox_counter = 1
        for pi, para in enumerate(doc.paragraphs, start=1):
            textboxes = self._extract_textboxes_from_paragraph(para)
            for tbx_local_index, box_info in enumerate(textboxes, start=1):
                for textbox_paragraph_index, textbox_text in enumerate(
                    box_info["paragraph_texts"],
                    start=1,
                ):
                    elements.append(ExtractedElement(
                        type="textbox",
                        element_id=(
                            f"textbox_{textbox_counter}_p{textbox_paragraph_index}"
                        ),
                        text=textbox_text,
                        empty=not bool(textbox_text.strip()),
                        part_name="word/document.xml",
                        container_path=(
                            f"/document/body/paragraph[{pi}]/"
                            f"textbox[{tbx_local_index}]/"
                            f"paragraph[{textbox_paragraph_index}]"
                        ),
                        textbox_index=textbox_counter,
                        paragraph_index=textbox_paragraph_index,
                        contains_linebreaks="\n" in textbox_text,
                        contains_tabs="\t" in textbox_text,
                        may_contain_special_runs=True,
                    ))
                textbox_counter += 1

        # 5. Footnotes
        for i, footnote_info in enumerate(self._extract_footnotes(), start=1):
            footnote_text = footnote_info["text"]
            xml_id = footnote_info["footnote_id"]

            elements.append(ExtractedElement(
                type="footnote",
                element_id=f"footnote_{i}",
                text=footnote_text.strip(),
                empty=not bool(footnote_text.strip()),
                part_name="word/footnotes.xml",
                container_path=f"/footnotes/footnote[@id='{xml_id}']",
                footnote_id=xml_id,
                contains_linebreaks="\n" in footnote_text,
                contains_tabs="\t" in footnote_text,
                may_contain_special_runs=True,
            ))

        structure["elements"] = [asdict(e) for e in elements]
        return structure

    def _extract_header_footer_elements(self) -> list[ExtractedElement]:
        elements: list[ExtractedElement] = []

        with ZipFile(self.filepath) as docx_zip:
            rels_root = etree.fromstring(
                docx_zip.read("word/_rels/document.xml.rels"),
                parser=self._xml_parser(),
            )
            document_root = etree.fromstring(
                docx_zip.read("word/document.xml"),
                parser=self._xml_parser(),
            )

            relationships = {}
            story_parts: dict[str, dict] = {}
            for rel in rels_root.findall(f"{{{PKG_REL_NS}}}Relationship"):
                rel_id = rel.get("Id")
                rel_type = rel.get("Type", "")
                if not rel_id or rel.get("TargetMode") == "External":
                    continue

                if rel_type.endswith("/header"):
                    kind = "header"
                elif rel_type.endswith("/footer"):
                    kind = "footer"
                else:
                    continue

                part_name = self._resolve_word_part(rel.get("Target", ""))
                relationships[rel_id] = (kind, part_name)
                story_parts.setdefault(part_name, {
                    "kind": kind,
                    "associations": [],
                })

            # Missing references inherit the corresponding story from the
            # previous section. Track default/first/even independently.
            effective: dict[tuple[str, str], str] = {}
            section_properties = document_root.xpath(".//w:sectPr", namespaces=NSMAP)
            for section_index, sect_pr in enumerate(section_properties, start=1):
                for kind in ("header", "footer"):
                    for ref in sect_pr.findall(f"{{{W_NS}}}{kind}Reference"):
                        rel_id = ref.get(f"{{{R_NS}}}id")
                        variant = ref.get(f"{{{W_NS}}}type", "default")
                        resolved = relationships.get(rel_id)
                        if resolved and resolved[0] == kind:
                            effective[(kind, variant)] = resolved[1]

                for (kind, variant), part_name in effective.items():
                    association = {
                        "section_index": section_index,
                        "variant": variant,
                    }
                    associations = story_parts[part_name]["associations"]
                    if association not in associations:
                        associations.append(association)

            for part_name, metadata in sorted(story_parts.items()):
                if part_name not in docx_zip.namelist():
                    self.logger.warning(
                        f"Header/footer relationship points to missing part: {part_name}"
                    )
                    continue

                root = etree.fromstring(
                    docx_zip.read(part_name),
                    parser=self._xml_parser(),
                )
                paragraphs = root.xpath(".//w:p", namespaces=NSMAP)
                kind = metadata["kind"]
                associations = metadata["associations"]
                section_indices = sorted({
                    item["section_index"] for item in associations
                })
                variants = sorted({item["variant"] for item in associations})
                part_token = self._story_part_token(part_name, kind)

                for paragraph_index, paragraph in enumerate(paragraphs, start=1):
                    text = self._paragraph_visible_text(paragraph)
                    variant = variants[0] if len(variants) == 1 else None
                    elements.append(ExtractedElement(
                        type=kind,
                        element_id=f"{kind}_{part_token}_p{paragraph_index}",
                        text=text,
                        empty=not bool(text.strip()),
                        part_name=part_name,
                        container_path=f"(.//w:p)[{paragraph_index}]",
                        section_index=section_indices[0] if section_indices else None,
                        header_index=paragraph_index if kind == "header" else None,
                        footer_index=paragraph_index if kind == "footer" else None,
                        paragraph_index=paragraph_index,
                        story_variant=variant,
                        section_indices=section_indices,
                        contains_linebreaks="\n" in text,
                        contains_tabs="\t" in text,
                        may_contain_special_runs=self._xml_paragraph_has_special_runs(paragraph),
                    ))

        return elements

    @staticmethod
    def _resolve_word_part(target: str) -> str:
        target = target.replace("\\", "/")
        if target.startswith("/"):
            resolved = posixpath.normpath(target.lstrip("/"))
        else:
            resolved = posixpath.normpath(posixpath.join("word", target))
        if resolved == ".." or resolved.startswith("../"):
            raise ValueError(f"Relationship target escapes DOCX package: {target}")
        return resolved

    @staticmethod
    def _story_part_token(part_name: str, kind: str) -> str:
        stem = os.path.splitext(posixpath.basename(part_name))[0]
        suffix = stem[len(kind):] if stem.startswith(kind) else stem
        suffix = re.sub(r"[^A-Za-z0-9]+", "_", suffix).strip("_") or "part"
        return suffix

    @staticmethod
    def _paragraph_visible_text(paragraph: etree._Element) -> str:
        parts = []
        for run in paragraph.iterdescendants(f"{{{W_NS}}}r"):
            nearest_paragraph = next(
                run.iterancestors(tag=f"{{{W_NS}}}p"),
                None,
            )
            if nearest_paragraph is not paragraph:
                continue
            for child in run:
                if child.tag == f"{{{W_NS}}}t":
                    parts.append(child.text or "")
                elif child.tag == f"{{{W_NS}}}tab":
                    parts.append("\t")
                elif child.tag in {f"{{{W_NS}}}br", f"{{{W_NS}}}cr"}:
                    parts.append("\n")
        return "".join(parts)

    @staticmethod
    def _xml_paragraph_has_special_runs(paragraph: etree._Element) -> bool:
        special_tags = {
            f"{{{W_NS}}}fldChar",
            f"{{{W_NS}}}instrText",
            f"{{{W_NS}}}drawing",
            f"{{{W_NS}}}footnoteReference",
            f"{{{W_NS}}}commentReference",
        }
        return any(node.tag in special_tags for node in paragraph.iterdescendants())

    @staticmethod
    def _xml_parser() -> etree.XMLParser:
        return etree.XMLParser(
            remove_blank_text=False,
            resolve_entities=False,
            no_network=True,
        )

    def _paragraph_may_contain_special_runs(self, para) -> bool:
        try:
            xml = para._element
            if xml.find(".//w:footnoteReference", namespaces=NSMAP) is not None:
                return True
            if xml.find(".//w:commentReference", namespaces=NSMAP) is not None:
                return True
            if xml.find(".//w:drawing", namespaces=NSMAP) is not None:
                return True
            return False
        except Exception:
            return True

    def _extract_textboxes_from_paragraph(self, paragraph):
        """
        Returnerar textboxes i ett stycke med text + xml-referens.
        """
        textboxes = []

        drawing_elements = paragraph._element.findall(
            ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
        )

        for drawing in drawing_elements:
            textbox_contents = drawing.findall(f".//{{{W_NS}}}txbxContent")
            paragraph_texts = []
            for textbox_content in textbox_contents:
                for textbox_paragraph in textbox_content.findall(f".//{{{W_NS}}}p"):
                    paragraph_texts.append(
                        self._visible_text_from_xml_paragraph(textbox_paragraph)
                    )

            if paragraph_texts:
                textboxes.append({
                    "xml": drawing,
                    "text": "".join(paragraph_texts),
                    "paragraph_texts": paragraph_texts,
                })

        return textboxes

    @staticmethod
    def _visible_text_from_xml_paragraph(paragraph: etree._Element) -> str:
        parts = []
        # python-docx BaseOxmlElement.xpath already supplies the standard
        # namespace mapping and does not accept lxml's namespaces argument.
        for node in paragraph.xpath(".//w:t | .//w:tab | .//w:br | .//w:cr"):
            if node.tag == f"{{{W_NS}}}t":
                parts.append(node.text or "")
            elif node.tag == f"{{{W_NS}}}tab":
                parts.append("\t")
            else:
                parts.append("\n")
        return "".join(parts)

    def _extract_footnotes(self):
        footnotes = []

        try:
            with ZipFile(self.filepath) as docx_zip:
                if "word/footnotes.xml" not in docx_zip.namelist():
                    return footnotes

                footnotes_xml = docx_zip.read("word/footnotes.xml")
                tree = etree.fromstring(footnotes_xml)

                for footnote in tree.findall("w:footnote", NSMAP):
                    footnote_id = footnote.get(f"{{{W_NS}}}id")

                    # hoppa över separatorer
                    if footnote_id in ("-1", "0"):
                        continue

                    texts = footnote.findall(".//w:t", NSMAP)
                    full_text = "".join([t.text for t in texts if t.text])

                    if full_text.strip():
                        footnotes.append({
                            "footnote_id": footnote_id,
                            "text": full_text,
                        })

        except Exception as e:
            self.logger.warning(f"Could not extract footnotes: {e}")

        return footnotes

    # ------------------------------------------------------------------
    # Hjälpmetoder för fortsatt pipeline
    # ------------------------------------------------------------------

    @staticmethod
    def _extract_docx_elements(filepath):
        """
        Behålls tills vidare för bakåtkompatibilitet med äldre editorer.
        """
        elements = {}
        doc = docx.Document(filepath)

        for i, para in enumerate(doc.paragraphs, start=1):
            elements[f"paragraph_{i}"] = para

        for ti, table in enumerate(doc.tables, start=1):
            for ri, row in enumerate(table.rows, start=1):
                for ci, cell in enumerate(row.cells, start=1):
                    elements[f"table_{ti}_cell_{ri}_{ci}"] = cell

        for si, section in enumerate(doc.sections, start=1):
            if hasattr(section.header, "paragraphs"):
                for hi, para in enumerate(section.header.paragraphs, start=1):
                    elements[f"header_s{si}_{hi}"] = para

            if hasattr(section.footer, "paragraphs"):
                for fi, para in enumerate(section.footer.paragraphs, start=1):
                    elements[f"footer_s{si}_{fi}"] = para

        textbox_counter = 1
        extractor = DocumentStructureExtractor(filepath, logger=_NullLogger())
        for para in doc.paragraphs:
            textboxes = extractor._extract_textboxes_from_paragraph(para)
            for box in textboxes:
                elements[f"textbox_{textbox_counter}"] = box["xml"]
                textbox_counter += 1

        for i, footnote in enumerate(extractor._extract_footnotes(), start=1):
            elements[f"footnote_{i}"] = footnote

        return doc, elements

    def save_as_json(self, output_path=None):
        if not self.structure:
            self.extract()
        if not output_path:
            output_path = self.filepath + "_structure.json"
        try:
            with open(output_path, "w", encoding="utf-8") as f:
                json.dump(self.structure, f, indent=2, ensure_ascii=False)
            return output_path
        except Exception as e:
            self.logger.error(f"Error saving JSON structure: {str(e)}")
            return None


class _NullLogger:
    def warning(self, msg): pass
    def error(self, msg): pass
    def info(self, msg): pass
    def debug(self, msg): pass


def main():
    import logging

    if len(sys.argv) != 2:
        print(f"Usage: python {os.path.basename(__file__)} <docx document>")
        sys.exit(1)

    filepath = sys.argv[1]

    logger = logging.getLogger("extractor-test")
    logger.setLevel(logging.INFO)
    handler = logging.StreamHandler(sys.stdout)
    handler.setFormatter(logging.Formatter('%(asctime)s - %(levelname)s - %(message)s'))
    logger.handlers.clear()
    logger.addHandler(handler)

    extractor = DocumentStructureExtractor(filepath, logger)
    extractor.extract()
    output_json = extractor.save_as_json()
    print(f"Structure saved to: {output_json}")


if __name__ == "__main__":
    main()
