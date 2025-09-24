import docx
from docx.shared import RGBColor
from docx.table import _Cell
from datetime import datetime, timezone
import os
import sys
import uuid
import shutil
import json
import fitz # PyMuPDF
import re
import zipfile
from lxml import etree
from tempfile import mkdtemp
import tempfile
from zipfile import ZipFile

from copy import deepcopy
from difflib import ndiff
from thefuzz import fuzz
try:
    from app.src.JBGDocumentStructureExtractor import DocumentStructureExtractor
except ModuleNotFoundError:
    from JBGDocumentStructureExtractor import DocumentStructureExtractor
    
# Settings
DEBUG = True
SUGGESTION = "Förslag"
MOTIVATION = "Motivering"
NSMAP = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
TEXT_SIM_SCORE_THRESHOLD = 90.0

class JBGDocumentEditor:
    
    def __init__(self, filepath, changes_path, include_motivations, logger):
        """
        :param filepath: Path to the input document (.docx or .pdf)
        :param changes_path: Path to the suggested changes JSON file
        """
        self.filepath = filepath
        self.changes = self._get_changes_from_json(changes_path)  
        self.logger = logger
        self.include_motivations = include_motivations
        self.footnote_changes = []
        self.textbox_changes = []
        self.ext = os.path.splitext(filepath)[1].lower()
        self.edited_document = None
        self.nsmap = NSMAP

    @staticmethod
    def _get_changes_from_json(json_filepath):
        with open(json_filepath, 'r', encoding='utf-8') as f:
            return json.load(f)    

    def apply_changes(self):
        try:
            # Choose method based on file type
            if self.filepath.endswith(".docx"):
                self._apply_changes_docx()
            elif self.filepath.endswith(".pdf"):
                self._apply_changes_pdf()
            else:
                raise ValueError("Unsupported file format. Use .docx or .pdf!")

        except Exception as e:
            self.logger.error(f"❌ Failed to apply changes: {e}")
            raise
        
    def _apply_changes_pdf(self):
        self.edited_document = self._annotate_pdf()
        
    def _apply_changes_docx(self):
        
        # Step 1: Apply visual markup
        markup_doc = self._edit_docx()

        # Step 2: Save intermediate version (used in both modes)
        intermediate_path = self._save_edited_document(doc=markup_doc, suffix="_intermediate")
        self.filepath = intermediate_path  # point everything to this file

        # Step 3: Patch footnotes.xml inside the ZIP
        if self.footnote_changes:
            self._edit_footnote_texts()

        # Steo 4: Patch textboxes in document.xml inside the ZIP
        if self.textbox_changes:
            self._edit_textbox_texts()
        
        self.edited_document = docx.Document(intermediate_path)
 
    def _save_edited_document(self, doc=None, suffix="_edited"):
        base, ext = os.path.splitext(self.filepath)
        output_path = f"{base}{suffix}{ext}"
        doc = doc or self.edited_document
        doc.save(output_path)
        self.logger.info(f"💾 Saved: {output_path}")
        return output_path

    def save_edited_document(self, output_path=None):

        if output_path is None:
            base, ext = os.path.splitext(self.filepath)
            suffix = "_edited" if ext.lower() == ".docx" else "_annotated"
            output_path = f"{base}{suffix}{ext}"

        if self.edited_document:
            try:
                self.edited_document.save(output_path)
                self.logger.info(f"💾 Saved document with visual edits or tracked changes to: {output_path}")
            except Exception as e:
                self.logger.error(f"❌ Failed to save edited document: {e}")
                raise
        else:
            try:
                if self.filepath != output_path:
                    shutil.copyfile(self.filepath, output_path)
                    self.logger.info(f"📝 Returned already-saved document with comments or annotations: {output_path}")
            except Exception as e:
                self.logger.error(f"❌ Failed to copy saved file: {e}")
                raise

        return output_path

    def _edit_docx(self):
        doc, elements = DocumentStructureExtractor._extract_docx_elements(self.filepath)
        applied_changes = set()

        for change in self.changes:
            element_id = change.get("element_id")
            old = change.get("old")
            new = change.get("new")
            motivation = change.get("motivation")

            if not element_id or element_id not in elements:
                self.logger.warning(f"⚠️ Skipping change — unknown or missing element_id: {element_id}")
                continue

            element = elements[element_id]
            if DEBUG: self.logger.debug(f" -- Considering element: {str(element)} with id {element_id}")

            # Standard Word elements with .text and .runs
            if hasattr(element, "text") and hasattr(element, "runs"):
                if DEBUG: self.logger.debug(f" -- Element has text and runs attribute")
                original_text = element.text  or ""
                if not original_text and isinstance(element, etree._Element):
                    original_text = self._get_joined_text_from_xml_element(element)
                if DEBUG: self.logger.debug(f" -- Original text was: {original_text}")

                normalized_old = self._normalize_text(old)
                normalized_origin = self._normalize_text(original_text)
                if normalized_old not in normalized_origin:
                    self.logger.debug(f"⚠️ Not identical match for '{old}' and '{original_text}' in {element_id}")
                    sim_score = fuzz.ratio(normalized_old, normalized_origin)
                    self.logger.debug(f"⚠️ Comparison similarity score: {sim_score} %")
                    if float(sim_score) < TEXT_SIM_SCORE_THRESHOLD:
                        self.logger.error(f"❌ No enough match for '{old}' in {element_id} (fuzzy match similarity score: {sim_score}")
                        continue
                
                # Build diff
                # Build a diff across the entire paragraph:
                # baseline is the original paragraph text; updated is baseline with the FIRST
                # occurrence of 'old' replaced by 'new' (with a whitespace-tolerant fallback).
                baseline_text = original_text
                updated_text = baseline_text
                idx = baseline_text.find(old)
                if idx != -1:
                    updated_text = baseline_text[:idx] + (new or "") + baseline_text[idx + len(old):]
                else:
                    try:
                        pattern = re.compile(r"\s+".join(map(re.escape, (old or "").split())), flags=re.UNICODE)
                        updated_text, n = pattern.subn(new or "", baseline_text, count=1)
                        if n == 0:
                            # fallback keeps behavior pre-patch
                            updated_text = new or ""
                    except Exception:
                        updated_text = baseline_text.replace(old or "", new or "", 1)
                diffed = self._diff_words(baseline_text, updated_text)
                # Capture footnoteReference runs together with their character positions
                # measured against the plain paragraph text (without refs).
                footnote_refs = []   # list[(offset_chars, run_xml)]
                pos_cursor = 0
                for run in element.runs:
                    xml_run = run._element
                    run_text = run.text or ""
                    if xml_run.find("w:footnoteReference", namespaces=self.nsmap) is not None:
                        footnote_refs.append((pos_cursor, deepcopy(xml_run)))
                    # Only text contributes to the visible paragraph text length
                    pos_cursor += len(run_text)

                # Clear original runs before rebuilding with markup.
                # (We re-insert any captured footnoteReference elements at their original offsets.)
                for _ in range(len(element.runs)):
                    element._element.remove(element.runs[0]._element)
                
                # Add formatted runs while re-inserting refs at their original baseline offsets.
                # Keep a baseline counter (characters that existed in the original paragraph).
                baseline_emitted = 0
                ref_i = 0
                total_refs = len(footnote_refs)
                # Insert any refs recorded at offset 0 (before any characters)
                while ref_i < total_refs and footnote_refs[ref_i][0] <= baseline_emitted:
                    element._element.append(footnote_refs[ref_i][1]); ref_i += 1

                for kind, val in diffed:
                    run = element.add_run(val)
                    if kind == "strike":
                        run.font.strike = True
                        run.font.color.rgb = RGBColor(255, 0, 0)
                    elif kind == "insert":
                        run.font.color.rgb = RGBColor(0, 128, 0)
                    # Advance baseline counter only for text that existed in the baseline
                    if kind in ("text", "strike"):
                        baseline_emitted += len(val)
                    # Insert any refs whose baseline offset is now reached
                    while ref_i < total_refs and footnote_refs[ref_i][0] <= baseline_emitted:
                        element._element.append(footnote_refs[ref_i][1])
                        ref_i += 1

                # Append any refs that might remain
                while ref_i < total_refs:
                    element._element.append(footnote_refs[ref_i][1])
                    ref_i += 1
                
                # Add motivation comment
                if self.include_motivations and motivation:
                    if "footer" in element_id or "header" in element_id:
                        self.logger.warning(f"⚠️ Skipping comment for {element_id} (unsupported in header/footer)")
                    else:
                        try:
                            element.add_comment(
                                text=motivation,
                                author="JBG klarspråkningstjänst",
                                initials="JBG"
                            )
                        except Exception as e:
                            self.logger.warning(f"⚠️ Could not add comment to {element_id}: {e}")

                applied_changes.add(element_id)
                self.logger.info(f"✅ Applied: '{old}' → '{new}' in {element_id}")
                continue
            
            # Handle table cells
            elif isinstance(element, _Cell):
                if DEBUG:
                    self.logger.debug(f" -- Element is _Cell")
                cell_handled = False

                normalized_old = self._normalize_text(old)

                # Step 1: Try exact substring match per paragraph
                for para in element.paragraphs:
                    para_text = self._normalize_text(para.text or "")
                    if normalized_old and normalized_old in para_text:
                        diffed = self._diff_words(old, new)

                        # Clear this paragraph’s runs
                        for _ in range(len(para.runs)):
                            para._element.remove(para.runs[0]._element)

                        # Rebuild with diff highlighting
                        for kind, val in diffed:
                            run = para.add_run(val)
                            if kind == "strike":
                                run.font.strike = True
                                run.font.color.rgb = RGBColor(255, 0, 0)
                            elif kind == "insert":
                                run.font.color.rgb = RGBColor(0, 128, 0)

                        if self.include_motivations and motivation:
                            try:
                                para.add_comment(
                                    text=motivation,
                                    author="JBG klarspråkningstjänst",
                                    initials="JBG",
                                )
                            except Exception as e:
                                self.logger.warning(
                                    f"⚠️ Could not add comment in table cell {element_id}: {e}"
                                )

                        applied_changes.add(element_id)
                        self.logger.info(f"✅ Applied in table cell {element_id} (exact match)")
                        cell_handled = True
                        break  # stop after first successful match

                # Step 2: Fallback fuzzy match if no exact hit
                if not cell_handled and element.paragraphs:
                    normalized_cell = self._normalize_text(
                        "\n".join(p.text.strip() for p in element.paragraphs)
                    )
                    sim_score = fuzz.ratio(normalized_old, normalized_cell)
                    self.logger.debug(f"⚠️ Fallback similarity score: {sim_score} %")

                    # Guard: avoid overwriting short/numeric content
                    if normalized_old.isdigit() or len(normalized_old) < 4:
                        self.logger.warning(
                            f"❌ Ignoring short/numeric '{old}' in {element_id} (avoiding false match)"
                        )
                    elif sim_score >= TEXT_SIM_SCORE_THRESHOLD:
                        para = element.paragraphs[0]
                        diffed = self._diff_words(old, new)
                        for _ in range(len(para.runs)):
                            para._element.remove(para.runs[0]._element)
                        for kind, val in diffed:
                            run = para.add_run(val)
                            if kind == "strike":
                                run.font.strike = True
                                run.font.color.rgb = RGBColor(255, 0, 0)
                            elif kind == "insert":
                                run.font.color.rgb = RGBColor(0, 128, 0)

                        if self.include_motivations and motivation:
                            try:
                                para.add_comment(
                                    text=motivation,
                                    author="JBG klarspråkningstjänst",
                                    initials="JBG",
                                )
                            except Exception as e:
                                self.logger.warning(
                                    f"⚠️ Could not add comment in table cell {element_id}: {e}"
                                )

                        applied_changes.add(element_id)
                        self.logger.info(f"✅ Applied in table cell {element_id} (fuzzy match)")
                        cell_handled = True

                # Step 3: Failure case
                if not cell_handled:
                    self.logger.error(f"❌ No match found in table cell {element_id}")
                    self.failed_changes.append(change)

                continue


            # Handle raw XML elements that are footnotes
            elif isinstance(element, etree._Element) and element_id.startswith("footnote"):
                if DEBUG: self.logger.debug("-- Element is raw XML and footnote")

                # Collect all footnote changes for later
                self.footnote_changes.append({
                    "element_id": element_id,
                    "footnote_id": change.get("footnote_id"),
                    "old": old,
                    "new": new,
                    "motivation": motivation
                })
                self.logger.info(f"Stored footnote '{element_id}' with new text '{new}' for later insertions")

            # Handle raw XML elements that are textboxes
            elif isinstance(element, etree._Element) and element_id.startswith("textbox"):
                if DEBUG: self.logger.debug("-- Element is raw XML and textbox")

                # Collect all footnote changes for later
                self.textbox_changes.append({
                    "element_id": element_id,
                    "textbox_id": change.get("textbox_id"),
                    "old": old,
                    "new": new,
                    "motivation": motivation
                })
                self.logger.info(f"Stored textbox '{element_id}' with new text '{new}' for later insertions")
            
            # Handle raw XML elements that are not footnotes nor textboxes(not supported)
            else:
                self.logger.warning(f"⚠️ Unsupported element type for {element_id}: {type(element)}")

        # Final reporting (mninus footnotes)
        for change in self.changes:
            element_id = change.get("element_id")
            if element_id not in applied_changes and not element_id.startswith(("footnote", "textbox")):
                self.logger.error(f"❌ Change not applied: {change}")

        return doc
    
    def _diff_words(self, old, new):
        """
        Returns a list of tuples: (type, text)
        type ∈ {"text", "insert", "strike"}
        """
        diff = list(ndiff(old.split(), new.split()))
        result = []

        for d in diff:
            tag, word = d[0], d[2:]
            if tag == " ":
                result.append(("text", word + " "))
            elif tag == "-":
                result.append(("strike", word + " "))
            elif tag == "+":
                result.append(("insert", word + " "))

        return result
    
    def _get_joined_text_from_xml_element(self, element):
        texts = element.findall(".//w:t", namespaces=self.nsmap)
        return "".join(t.text for t in texts if t.text)

    def _edit_footnote_texts(self):
        """
        Patchar word/footnotes.xml direkt i .docx (ZIP) och uppdaterar
        word/document.xml så att ev. motivations-kommentarer ankeras vid
        fotnotreferensen i brödtexten – placerad ett steg längre bort
        från superskriptet när möjligt.

        Strategi:
        - För varje fotnot-ID i self.footnote_changes:
            * Läs full baslinjetext i fotnoten
            * Bygg om fotnoten:
                - bevara ENDAST fotnotRef-run (själva markören)
                - lägg till ett rent mellanslag (syntetiskt), inte en hel run
                från originalet som råkar innehålla text
                - röd/överstruken run = hela gamla texten
                - grön run = hela nya texten
            * Om motivation finns och self.include_motivations är True:
                - säkerställ comments.xml, CT-override och rels
                - lägg in kommentar med author/initials
                - ankara <w:commentReference> efter fotnotsreferensen,
                men helst efter nästa whitespace-run (ett steg längre bort)
        """

        import shutil
        import zipfile
        from copy import deepcopy
        from datetime import datetime
        from lxml import etree

        # --- Namnrymder
        ns = dict(self.nsmap or {})
        W_NS   = ns.get("w", "http://schemas.openxmlformats.org/wordprocessingml/2006/main")
        REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
        ODT_REL_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        CT_NS  = "http://schemas.openxmlformats.org/package/2006/content-types"
        XML_NS = "http://www.w3.org/XML/1998/namespace"
        ns.setdefault("w", W_NS)

        def w(tag): return f"{{{W_NS}}}{tag}"

        # --- Hjälpare för ZIP/bytes ---
        def _read_member(zf, name):
            try:
                with zf.open(name) as f:
                    return f.read()
            except KeyError:
                return None

        def _rewrite_zip(src_zip_path, overrides: dict):
            tmp_out = src_zip_path + ".tmp_edit_footnotes"
            with zipfile.ZipFile(src_zip_path, "r") as zin, \
                zipfile.ZipFile(tmp_out, "w", compression=zipfile.ZIP_DEFLATED) as zout:
                names = set(zin.namelist())
                for name in names:
                    if name in overrides:
                        continue
                    zout.writestr(name, zin.read(name))
                for name, data in overrides.items():
                    zout.writestr(name, data)
            shutil.move(tmp_out, src_zip_path)

        if not getattr(self, "footnote_changes", None):
            self.logger.info("No footnote changes to apply.")
            return

        # --- Läs in nödvändiga delar
        with zipfile.ZipFile(self.filepath, "r") as zin:
            footnotes_xml = _read_member(zin, "word/footnotes.xml")
            document_xml  = _read_member(zin, "word/document.xml")
            rels_xml      = _read_member(zin, "word/_rels/document.xml.rels")
            ctypes_xml    = _read_member(zin, "[Content_Types].xml")
            comments_xml  = _read_member(zin, "word/comments.xml")

        if footnotes_xml is None:
            self.logger.warning("word/footnotes.xml saknas – inga fotnoter att patcha.")
            return
        if document_xml is None or ctypes_xml is None:
            self.logger.error("Nödvändiga delar saknas (document.xml eller [Content_Types].xml).")
            return
        if rels_xml is None:
            # skapa minimalt rels-dokument om det saknas
            rels_root = etree.Element(f"{{{REL_NS}}}Relationships", nsmap={None: REL_NS})
            rels_xml = etree.tostring(etree.ElementTree(rels_root), xml_declaration=True, encoding="utf-8", standalone="yes")

        parser = etree.XMLParser(remove_blank_text=False, ns_clean=True, recover=True)
        foot_tree = etree.fromstring(footnotes_xml, parser=parser)
        doc_tree  = etree.fromstring(document_xml,  parser=parser)
        rels_tree = etree.fromstring(rels_xml,      parser=parser)
        ct_tree   = etree.fromstring(ctypes_xml,    parser=parser)
        com_tree  = etree.fromstring(comments_xml,  parser=parser) if comments_xml is not None else None

        # --- Comments-stöd ---
        def _ensure_comments_part():
            nonlocal com_tree, rels_tree, ct_tree
            changed_rels = False
            changed_ct   = False

            if com_tree is None:
                com_tree = etree.Element(w("comments"), nsmap={"w": W_NS})

            # Content_Types override
            has_override = any(
                (ov.get("PartName") == "/word/comments.xml")
                for ov in ct_tree.findall(f"{{{CT_NS}}}Override")
            )
            if not has_override:
                new_ov = etree.SubElement(ct_tree, f"{{{CT_NS}}}Override")
                new_ov.set("PartName", "/word/comments.xml")
                new_ov.set("ContentType", "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml")
                changed_ct = True

            # Relationship till comments.xml
            has_rel = any(
                (rel.get("Type") == f"{ODT_REL_NS}/comments" and rel.get("Target") == "comments.xml")
                for rel in rels_tree.findall(f"{{{REL_NS}}}Relationship")
            )
            if not has_rel:
                existing_ids = {rel.get("Id") for rel in rels_tree.findall(f"{{{REL_NS}}}Relationship")}
                i = 1
                while f"rId{i}" in existing_ids:
                    i += 1
                new_rel = etree.SubElement(rels_tree, f"{{{REL_NS}}}Relationship")
                new_rel.set("Id", f"rId{i}")
                new_rel.set("Type", f"{ODT_REL_NS}/comments")
                new_rel.set("Target", "comments.xml")
                changed_rels = True

            # Nästa comment-id
            max_id = -1
            for c in com_tree.findall(f".//{w('comment')}"):
                try:
                    cid = int(c.get(f"{{{W_NS}}}id"))
                    if cid > max_id:
                        max_id = cid
                except Exception:
                    pass
            next_id = max_id + 1
            return com_tree, next_id, changed_rels, changed_ct

        def _add_comment(comment_id: int, text: str):
            c = etree.SubElement(com_tree, w("comment"))
            c.set(f"{{{W_NS}}}id", str(comment_id))
            # sätt avsändare/initialer och datum
            c.set(f"{{{W_NS}}}author", "JBG klarspråkningstjänst")
            c.set(f"{{{W_NS}}}initials", "JBG")
            c.set(f"{{{W_NS}}}date", datetime.utcnow().strftime("%Y-%m-%dT%H:%M:%SZ"))

            p = etree.SubElement(c, w("p"))
            r = etree.SubElement(p, w("r"))
            t = etree.SubElement(r, w("t"))
            t.set(f"{{{XML_NS}}}space", "preserve")
            t.text = text or ""

        def _is_whitespace_run(r):
            t = r.find(w("t"))
            if t is None or t.text is None:
                return False
            return t.text.strip("") == "" and len(t.text) > 0  # innehåller bara whitespace

        def _insert_comment_reference_after_ref(footnote_id: str, comment_id: int):
            """
            Placera commentReference ett steg bort:
            - om nästa syskon-run efter referensen är en ren whitespace-run,
                lägg commentReference EFTER den runnen
            - annars, om det finns en efterföljande run, lägg efter den
            - annars efter referens-runnen
            """
            ref = doc_tree.find(f".//{w('footnoteReference')}[@{w('id')}='{footnote_id}']")
            if ref is None:
                self.logger.warning(f"Could not find footnoteReference id={footnote_id} in document.xml")
                return False

            # referensen ska ligga i en w:r
            parent_r = ref.getparent()
            if parent_r is None or parent_r.tag != w("r"):
                parent = ref.getparent()
                idx = parent.index(ref)
                crun = etree.Element(w("r"))
                etree.SubElement(crun, w("commentReference")).set(f"{{{W_NS}}}id", str(comment_id))
                parent.insert(idx + 1, crun)
                return True

            p = parent_r.getparent()
            if p is None:
                self.logger.warning("Unexpected XML around footnoteReference.")
                return False

            # hitta infogningsindex
            insert_after_index = p.index(parent_r)

            # om nästa syskon-run är whitespace, hoppa över den
            if insert_after_index + 1 < len(p):
                next_node = p[insert_after_index + 1]
                if next_node.tag == w("r") and _is_whitespace_run(next_node):
                    insert_after_index += 1
                elif next_node.tag == w("r"):
                    # om det inte är whitespace men finns – lägg efter den (ett steg längre bort)
                    insert_after_index += 1

            crun = etree.Element(w("r"))
            etree.SubElement(crun, w("commentReference")).set(f"{{{W_NS}}}id", str(comment_id))
            p.insert(insert_after_index + 1, crun)
            return True

        # --- Patcha fotnoter
        changes_applied = 0
        for ch in (self.footnote_changes or []):
            footnote_id = str(ch.get("footnote_id", "")).strip()
            old_text    = ch.get("old") or ""
            new_text    = ch.get("new") or ""
            motivation  = ch.get("motivation")

            if not footnote_id:
                self.logger.warning("Skipping a footnote change without footnote_id.")
                continue

            fn = foot_tree.find(f".//{w('footnote')}[@{w('id')}='{footnote_id}']")
            if fn is None:
                self.logger.warning(f"Footnote id={footnote_id} not found in footnotes.xml")
                continue

            # Full baslinjetext (alla w:t)
            baseline = "".join(t.text or "" for t in fn.findall(f".//{w('t')}"))

            # Bevara endast fotnotRef-run från första paragrafen
            first_p = fn.find(f".//{w('p')}")
            saved_ppr = deepcopy(first_p.find(w("pPr"))) if first_p is not None else None

            preserved_ref_run = None
            if first_p is not None:
                for r in first_p.findall(w("r")):
                    if r.find(w("footnoteRef")) is not None:
                        preserved_ref_run = deepcopy(r)
                        break

            # Töm hela fotnoten och bygg en (1) ny paragraf
            for child in list(fn):
                fn.remove(child)
            new_p = etree.SubElement(fn, w("p"))
            if saved_ppr is not None:
                new_p.append(saved_ppr)

            # Lägg tillbaka bara referens-runnen
            if preserved_ref_run is not None:
                new_p.append(preserved_ref_run)

            # Lägg till ett syntetiskt mellanslag (run med " ")
            r_space = etree.SubElement(new_p, w("r"))
            t_space = etree.SubElement(r_space, w("t"))
            t_space.set(f"{{{XML_NS}}}space", "preserve")
            t_space.text = " "

            # Rollback-1: hela gamla röd/strike + hela nya grön
            def _add_run(text, *, red_strike=False, green=False):
                r = etree.SubElement(new_p, w("r"))
                rpr = etree.SubElement(r, w("rPr"))
                if red_strike:
                    etree.SubElement(rpr, w("color"), {f"{{{W_NS}}}val": "FF0000"})
                    etree.SubElement(rpr, w("strike"))
                if green:
                    etree.SubElement(rpr, w("color"), {f"{{{W_NS}}}val": "008000"})
                t = etree.SubElement(r, w("t"))
                t.set(f"{{{XML_NS}}}space", "preserve")
                t.text = text if text is not None else ""

            _add_run(baseline, red_strike=True)
            _add_run(new_text, green=True)

            changes_applied += 1

            # Kommentar/motivation
            if self.include_motivations and motivation:
                com_root, next_cid, rels_changed, ct_changed = _ensure_comments_part()
                _add_comment(next_cid, str(motivation))
                if _insert_comment_reference_after_ref(footnote_id, next_cid):
                    self.logger.info(f"Anchored comment (id={next_cid}) near footnoteRef id={footnote_id}.")
                else:
                    self.logger.warning(f"Failed to anchor comment for footnoteRef id={footnote_id}.")

        if changes_applied == 0:
            self.logger.info("No matching footnotes were changed.")
            return

        # --- Skriv tillbaka
        footnotes_out = etree.tostring(foot_tree, xml_declaration=True, encoding="utf-8", standalone="yes")
        document_out  = etree.tostring(doc_tree,  xml_declaration=True, encoding="utf-8", standalone="yes")
        rels_out      = etree.tostring(rels_tree, xml_declaration=True, encoding="utf-8", standalone="yes")
        ctypes_out    = etree.tostring(ct_tree,   xml_declaration=True, encoding="utf-8", standalone="yes")
        overrides = {
            "word/footnotes.xml": footnotes_out,
            "word/document.xml":  document_out,
            "word/_rels/document.xml.rels": rels_out,
            "[Content_Types].xml": ctypes_out,
        }
        if com_tree is not None:
            comments_out = etree.tostring(com_tree, xml_declaration=True, encoding="utf-8", standalone="yes")
            overrides["word/comments.xml"] = comments_out

        _rewrite_zip(self.filepath, overrides)
        self.logger.info(f"✅ Applied {changes_applied} footnote change(s) with safe whitespace handling and improved comment anchoring.")

    def _edit_textbox_texts(self):
        """
        Patchar word/document.xml för att uppdatera text i textboxes (<w:txbxContent>).
        Strategi:
        - För varje entry i self.textbox_changes:
            * Loopa över alla textboxes i dokumentet
            * Läs deras text och jämför mot 'old'
            * Om match: bygg om runs i den textboxen med diff (strike/red och insert/green)
        """
        import shutil, zipfile
        from lxml import etree

        if not getattr(self, "textbox_changes", None):
            self.logger.info("No textbox changes to apply.")
            return

        ns = dict(self.nsmap or {})
        W_NS = ns.get("w", "http://schemas.openxmlformats.org/wordprocessingml/2006/main")
        ns.setdefault("w", W_NS)

        def w(tag): 
            return f"{{{W_NS}}}{tag}"

        def _add_run(parent, text, *, red_strike=False, green=False):
            """Helper: create <w:r><w:t> with optional formatting"""
            r = etree.SubElement(parent, w("r"))
            rpr = etree.SubElement(r, w("rPr"))
            if red_strike:
                etree.SubElement(rpr, w("color"), {f"{{{W_NS}}}val": "FF0000"})
                etree.SubElement(rpr, w("strike"))
            if green:
                etree.SubElement(rpr, w("color"), {f"{{{W_NS}}}val": "008000"})
            t = etree.SubElement(r, w("t"))
            XML_NS = "http://www.w3.org/XML/1998/namespace"
            t.set(f"{{{XML_NS}}}space", "preserve")
            t.text = text if text is not None else ""


        # --- Läs och patcha document.xml
        with zipfile.ZipFile(self.filepath, "r") as zin:
            document_xml = zin.read("word/document.xml")

        parser = etree.XMLParser(remove_blank_text=False, ns_clean=True, recover=True)
        doc_tree = etree.fromstring(document_xml, parser=parser)

        changes_applied = 0
        for ch in (self.textbox_changes or []):
            old_text = ch.get("old") or ""
            new_text = ch.get("new") or ""
            motivation = ch.get("motivation")
            element_id = ch.get("element_id")

            norm_old = self._normalize_textbox_text(old_text)
            if not norm_old:
                continue

            # Iterate over all textboxes in the document
            found_box = False
            for txbx in doc_tree.findall(".//w:txbxContent", namespaces=ns):
                local_text_nodes = txbx.findall(".//w:t", namespaces=ns)
                local_full_text = "".join(t.text or "" for t in local_text_nodes)
                norm_local = self._normalize_textbox_text(local_full_text)

                self.logger.debug(f"[Textbox {element_id}] normalized_old='{norm_old}', normalized_local='{norm_local}'"
)
                if norm_old not in norm_local:
                    continue  # not this box, check next

                # Found the correct textbox
                diffed = self._diff_words(old_text, new_text)

                # Clear old runs
                for t in local_text_nodes:
                    parent_r = t.getparent()
                    if parent_r is not None:
                        parent_p = parent_r.getparent()
                        if parent_p is not None:
                            parent_p.remove(parent_r)

                # Ensure at least one paragraph
                p = txbx.find("w:p", namespaces=ns)
                if p is None:
                    p = etree.SubElement(txbx, w("p"))

                # Insert diffed runs
                # enforce consistent order: strike (old) before insert (new)
                buffered = []
                for kind, val in diffed:
                    if kind == "text":
                        # flush any buffered strike+insert before continuing
                        for bkind, bval in buffered:
                            if bkind == "strike":
                                _add_run(p, bval, red_strike=True)
                            elif bkind == "insert":
                                _add_run(p, bval, green=True)
                        buffered.clear()
                        _add_run(p, val)
                    else:
                        buffered.append((kind, val))

                # flush any remaining at end
                for bkind, bval in buffered:
                    if bkind == "strike":
                        _add_run(p, bval, red_strike=True)
                    elif bkind == "insert":
                        _add_run(p, bval, green=True)

                if self.include_motivations and motivation:
                    self.logger.info(f"ℹ️ Skipping comment insertion in textbox {element_id} (not implemented)")

                self.logger.info(f"✅ Applied in textbox {element_id}")
                changes_applied += 1
                found_box = True
                break  # stop after first matching textbox

            if not found_box:
                self.logger.warning(f"❌ No match for textbox {element_id}")

        if changes_applied == 0:
            self.logger.info("No matching textboxes were changed.")
            return

        # --- Skriv tillbaka document.xml
        new_doc_xml = etree.tostring(doc_tree, xml_declaration=True, encoding="utf-8", standalone="yes")

        tmp_out = self.filepath + ".tmp_edit_textboxes"
        with zipfile.ZipFile(self.filepath, "r") as zin, \
            zipfile.ZipFile(tmp_out, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                if item.filename == "word/document.xml":
                    zout.writestr(item, new_doc_xml)
                else:
                    zout.writestr(item, zin.read(item.filename))

        shutil.move(tmp_out, self.filepath)
        self.logger.info(f"✅ Applied {changes_applied} textbox change(s).")

    
    @staticmethod
    def _normalize_text(text):
        # Replace all whitespace (tabs, newlines, etc.) with single spaces
        return re.sub(r'\s+', ' ', text).strip()
    
    @staticmethod
    def _normalize_textbox_text(text: str) -> str:
            if not text:
                return ""
            # Remove all whitespace characters
            s = re.sub(r"\s+", "", text)
            # Normalize dashes
            s = s.replace("–", "-").replace("—", "-")
            return s.strip()
    
    @staticmethod
    def _clean_pdf_text(text):

        # Remove digits stuck to end of words: deltidsarbete15 → deltidsarbete
        text = re.sub(r"(?<=[a-zA-Z])\d{1,3}\b", "", text)

        # Remove digits at end of word before punctuation: deltidsarbete15. → deltidsarbete.
        text = re.sub(r"(?<=[a-zA-Z])\d{1,3}(?=[.,])", "", text)

        return text.strip()
    
    def _annotate_pdf(self):
        doc = fitz.open(self.filepath)

        for change in self.changes:
            if "page" not in change or "old" not in change:
                continue

            page_index = change["page"] - 1
            if page_index >= len(doc):
                continue

            old = change["old"]
            new = change.get("new", "")
            motivation = change.get("motivation", "")
            target_line = change.get("line")
            page = doc[page_index]

            # Get all text blocks on page, sorted top→down.
            # PyMuPDF coordinates increase downward, so smaller y means closer to the top.
            blocks = sorted(page.get_text("blocks"), key=lambda b: b[1])  # y1 ascending (top→down)
            line_texts = [(i + 1, b, b[4].strip()) for i, b in enumerate(blocks) if b[4].strip()]

            match_found = False

            # Try target line first
            if target_line:
                for line_no, block, text in line_texts:
                    if line_no == target_line and (self._normalize_text(old) in self._normalize_text(self._clean_pdf_text(text)) 
                                                   or self._normalize_text(self._clean_pdf_text(text)) in self._normalize_text(self._clean_pdf_text(old))):
                        rects = page.search_for(old, clip=fitz.Rect(block[:4]))
                        for rect in rects:
                            highlight = page.add_highlight_annot(rect)
                            if new:
                                if self.include_motivations:
                                    highlight.set_info(content=f"{SUGGESTION}: {new} \n\n{MOTIVATION}: {motivation}")
                                else:
                                    highlight.set_info(content=f"{SUGGESTION}: {new}")
                        match_found = True
                        self.logger.info(f"✅ Applied: '{old}' → '{new}' on page {page_index + 1}, line {target_line}")
                        break

                # If not found, try nearby lines
                if not match_found:
                    for line_no, block, text in line_texts:
                        if abs(line_no - target_line) <= 2 and (self._normalize_text(old) in self._normalize_text(self._clean_pdf_text(text)) 
                                                  or self._normalize_text(self._clean_pdf_text(text)) in self._normalize_text(self._clean_pdf_text(old))):
                            rects = page.search_for(old, clip=fitz.Rect(block[:4]))
                            for rect in rects:
                                highlight = page.add_highlight_annot(rect)
                                if new:
                                    if self.include_motivations:
                                        highlight.set_info(content=f"{SUGGESTION}: {new} \n\n{MOTIVATION}: {motivation}")
                                    else:
                                        highlight.set_info(content=f"{SUGGESTION}: {new}")
                            self.logger.warning(f"Could not find '{old}' on page {change['page']}, line {target_line} — did you mean line {line_no}?")
                            match_found = True
                            break
            else:
                # Fallback if no line provided
                rects = page.search_for(old)
                for rect in rects:
                    highlight = page.add_highlight_annot(rect)
                    if new:
                        if self.include_motivations:
                            highlight.set_info(content=f"{SUGGESTION}: {new} \n\n{MOTIVATION}: {motivation}")
                        else:
                            highlight.set_info(content=f"{SUGGESTION}: {new}")
                match_found = bool(rects)

            if not match_found:
                self.logger.error(f"❌No match found for '{old}' on page {change['page']}.")

        self._deduplicate_annotations(doc)
        return doc
    
    def _deduplicate_annotations(self, doc, distance_threshold=5):
        
        num_duplicates = 0
        for page in doc:
            existing = []
            to_remove = []

            annots = page.annots()
            if not annots:
                continue
            for annot in annots:
                if annot.type[0] != 8:  # 8 = Highlight, 1 = FreeText
                    continue

                rect = annot.rect
                content = (annot.info.get("content") or "").strip()

                is_duplicate = False
                for seen_rect, seen_text in existing:
                    if (
                        content == seen_text or
                        (abs(rect.x0 - seen_rect.x0) < distance_threshold and
                        abs(rect.y0 - seen_rect.y0) < distance_threshold and
                        abs(rect.x1 - seen_rect.x1) < distance_threshold and
                        abs(rect.y1 - seen_rect.y1) < distance_threshold)
                    ):
                        is_duplicate = True
                        break

                if is_duplicate:
                    to_remove.append(annot)
                else:
                    existing.append((rect, content))

            num_duplicates += len(to_remove)
            for annot in to_remove:
                page.delete_annot(annot)

        self.logger.info(f"Removed {num_duplicates} duplicate annotations.")
        return
