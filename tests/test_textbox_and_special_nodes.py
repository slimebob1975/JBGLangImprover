import logging
import tempfile
import unittest
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml import OxmlElement
from lxml import etree

from app.src.JBGChangePlanner import ChangePlanner
from app.src.JBGDocumentStructureExtractor import DocumentStructureExtractor
from app.src.JBGDocxPackage import W_NS, DocxPackage
from app.src.JBGLangImprovSuggestorAI import SuggestedChange
from app.src.JBGTokenDiffEngine import TokenDiffEngine
from app.src.JBGTrackedChangesRenderer import TrackedChangesRenderer


NS = {"w": W_NS}


class TextboxAndSpecialNodeTests(unittest.TestCase):
    def setUp(self):
        self.temp_dir = tempfile.TemporaryDirectory()
        self.root = Path(self.temp_dir.name)
        self.logger = logging.getLogger(f"textbox-special-test-{id(self)}")
        self.logger.handlers.clear()
        self.logger.addHandler(logging.NullHandler())

    def tearDown(self):
        self.temp_dir.cleanup()

    def test_textboxes_follow_drawing_order_and_extract_each_paragraph(self):
        source = self.root / "textboxes.docx"
        self._create_textbox_fixture(source)

        structure = DocumentStructureExtractor(str(source), self.logger).extract()
        textboxes = [
            element for element in structure["elements"]
            if element["type"] == "textbox"
        ]

        self.assertEqual(
            [element["element_id"] for element in textboxes],
            ["textbox_1_p1", "textbox_2_p1", "textbox_2_p2"],
        )
        self.assertEqual(
            [element["text"] for element in textboxes],
            ["Första rutan", "Andra rutans första stycke", "ha diagranm%"],
        )

    def test_tracked_changes_target_the_requested_textbox_paragraph(self):
        source = self.root / "textboxes.docx"
        output = self.root / "textboxes-tracked.docx"
        self._create_textbox_fixture(source)
        structure = DocumentStructureExtractor(str(source), self.logger).extract()
        suggestion = SuggestedChange(
            element_type="textbox",
            element_id="textbox_2_p2",
            footnote_id=None,
            old="ha diagranm%",
            new="(se diagram)",
            motivation="Regressionstest",
            match_status="exact",
        )

        results = self._render(source, output, structure, suggestion)

        self.assertTrue(results[0].applied, results[0].message)
        paragraphs = self._drawing_textbox_paragraphs(output, drawing_index=2)
        self.assertEqual(self._all_text(paragraphs[0]), "Andra rutans första stycke")
        self.assertIn("ha diagranm%", self._all_text(paragraphs[1]))
        self.assertIn("(se diagram)", self._all_text(paragraphs[1]))
        self.assertTrue(paragraphs[1].xpath(".//w:del", namespaces=NS))
        self.assertTrue(paragraphs[1].xpath(".//w:ins", namespaces=NS))

    def test_legacy_textbox_id_resolves_to_first_paragraph_of_the_drawing(self):
        source = self.root / "textboxes.docx"
        output = self.root / "textbox-legacy.docx"
        self._create_textbox_fixture(source)
        structure = DocumentStructureExtractor(str(source), self.logger).extract()
        suggestion = SuggestedChange(
            element_type="textbox",
            element_id="textbox_2",
            footnote_id=None,
            old="Andra rutans första stycke",
            new="Andra rutans inledande stycke",
            motivation="Bakåtkompatibilitet",
            match_status="exact",
        )

        results = self._render(source, output, structure, suggestion)

        self.assertTrue(results[0].applied, results[0].message)
        first = self._drawing_textbox_paragraphs(output, drawing_index=2)[0]
        self.assertIn("Andra rutans inledande stycke", self._all_text(first))

    def test_tracked_change_can_end_at_a_linebreak_in_the_same_run(self):
        source = self.root / "linebreak.docx"
        output = self.root / "linebreak-tracked.docx"
        document = Document()
        paragraph = document.add_paragraph("Inledning ")
        run = paragraph.add_run("måltext")
        run.add_break()
        run.add_text("Efter")
        document.save(source)

        structure = DocumentStructureExtractor(str(source), self.logger).extract()
        suggestion = SuggestedChange(
            element_type="paragraph",
            element_id="paragraph_1",
            footnote_id=None,
            old="måltext\n",
            new="ersättning",
            motivation="Specialnodstest",
            match_status="exact",
        )

        results = self._render(source, output, structure, suggestion)

        self.assertTrue(results[0].applied, results[0].message)
        paragraph_xml = self._main_paragraph(output)
        self.assertTrue(paragraph_xml.xpath(".//w:del", namespaces=NS))
        self.assertTrue(paragraph_xml.xpath(".//w:ins", namespaces=NS))
        self.assertIn("Efter", self._all_text(paragraph_xml))

    def _render(self, source, output, structure, suggestion):
        plans = ChangePlanner(
            structure,
            TokenDiffEngine(logger=self.logger),
            self.logger,
        ).build_plans([suggestion])
        self.assertEqual(len(plans), 1)
        with DocxPackage(str(source), self.logger) as package:
            results = TrackedChangesRenderer(package, self.logger).apply_plans(plans)
            package.save(str(output))
        return results

    @staticmethod
    def _create_textbox_fixture(path: Path):
        document = Document()
        first = document.add_paragraph()
        TextboxAndSpecialNodeTests._append_fallback(first, "Första fallback")
        TextboxAndSpecialNodeTests._append_drawing(first, ["Första rutan"])

        second = document.add_paragraph()
        TextboxAndSpecialNodeTests._append_fallback(second, "Andra fallback")
        TextboxAndSpecialNodeTests._append_drawing(
            second,
            ["Andra rutans första stycke", "ha diagranm%"],
        )
        document.save(path)

    @staticmethod
    def _append_fallback(paragraph, text):
        pict = OxmlElement("w:pict")
        content = OxmlElement("w:txbxContent")
        content.append(TextboxAndSpecialNodeTests._make_xml_paragraph(text))
        pict.append(content)
        paragraph._p.append(pict)

    @staticmethod
    def _append_drawing(paragraph, texts):
        drawing = OxmlElement("w:drawing")
        content = OxmlElement("w:txbxContent")
        for text in texts:
            content.append(TextboxAndSpecialNodeTests._make_xml_paragraph(text))
        drawing.append(content)
        paragraph._p.append(drawing)

    @staticmethod
    def _make_xml_paragraph(text):
        paragraph = OxmlElement("w:p")
        run = OxmlElement("w:r")
        text_element = OxmlElement("w:t")
        text_element.text = text
        run.append(text_element)
        paragraph.append(run)
        return paragraph

    @staticmethod
    def _drawing_textbox_paragraphs(path: Path, drawing_index: int):
        with ZipFile(path) as archive:
            root = etree.fromstring(archive.read("word/document.xml"))
        drawings = root.xpath("//w:body/w:p//w:drawing[.//w:txbxContent]", namespaces=NS)
        return drawings[drawing_index - 1].xpath(".//w:txbxContent//w:p", namespaces=NS)

    @staticmethod
    def _main_paragraph(path: Path):
        with ZipFile(path) as archive:
            root = etree.fromstring(archive.read("word/document.xml"))
        return root.xpath("//w:body/w:p", namespaces=NS)[0]

    @staticmethod
    def _all_text(element):
        return "".join(element.xpath(".//w:t/text() | .//w:delText/text()", namespaces=NS))


if __name__ == "__main__":
    unittest.main()
