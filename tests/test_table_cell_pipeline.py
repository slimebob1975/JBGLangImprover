import logging
import tempfile
import unittest
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from lxml import etree

from app.src.JBGChangePlanner import ChangePlanner
from app.src.JBGDocumentStructureExtractor import DocumentStructureExtractor
from app.src.JBGDocxPackage import W_NS, DocxPackage
from app.src.JBGLangImprovSuggestorAI import SuggestedChange
from app.src.JBGSimpleMarkupRenderer import SimpleMarkupRenderer
from app.src.JBGTokenDiffEngine import TokenDiffEngine
from app.src.JBGTrackedChangesRenderer import TrackedChangesRenderer


NS = {"w": W_NS}


class TableCellPipelineTests(unittest.TestCase):
    def setUp(self):
        self.temp_dir = tempfile.TemporaryDirectory()
        self.root = Path(self.temp_dir.name)
        self.logger = logging.getLogger(f"table-cell-test-{id(self)}")
        self.logger.handlers.clear()
        self.logger.addHandler(logging.NullHandler())
        self.source = self.root / "source.docx"

        document = Document()
        cell = document.add_table(rows=1, cols=1).cell(0, 0)
        cell.paragraphs[0].text = "Första stycket ska vara orört."
        cell.add_paragraph("Andra stycket innehåller måltexten.")
        document.save(self.source)

    def tearDown(self):
        self.temp_dir.cleanup()

    def test_extracts_each_table_cell_paragraph_as_a_separate_element(self):
        structure = DocumentStructureExtractor(str(self.source), self.logger).extract()
        cells = [element for element in structure["elements"] if element["type"] == "table_cell"]

        self.assertEqual(
            [element["element_id"] for element in cells],
            ["table_1_cell_1_1_p1", "table_1_cell_1_1_p2"],
        )
        self.assertEqual([element["paragraph_index"] for element in cells], [1, 2])
        self.assertEqual(cells[0]["text"], "Första stycket ska vara orört.")
        self.assertEqual(cells[1]["text"], "Andra stycket innehåller måltexten.")

    def test_simple_markup_targets_the_second_table_cell_paragraph(self):
        output = self.root / "simple.docx"
        results = self._render_second_paragraph(output, SimpleMarkupRenderer)

        self.assertTrue(all(result.applied for result in results), [
            result.message for result in results
        ])
        paragraphs = self._table_paragraphs(output)
        self.assertEqual(self._all_text(paragraphs[0]), "Första stycket ska vara orört.")
        self.assertIn("måltexten", self._all_text(paragraphs[1]))
        self.assertIn("den tydliga texten", self._all_text(paragraphs[1]))
        self.assertTrue(paragraphs[1].xpath(".//w:strike", namespaces=NS))

    def test_tracked_changes_target_the_second_table_cell_paragraph(self):
        output = self.root / "tracked.docx"
        results = self._render_second_paragraph(output, TrackedChangesRenderer)

        self.assertTrue(all(result.applied for result in results), [
            result.message for result in results
        ])
        paragraphs = self._table_paragraphs(output)
        self.assertEqual(self._all_text(paragraphs[0]), "Första stycket ska vara orört.")
        self.assertTrue(paragraphs[1].xpath(".//w:del", namespaces=NS))
        self.assertTrue(paragraphs[1].xpath(".//w:ins", namespaces=NS))
        self.assertIn("den tydliga texten", self._all_text(paragraphs[1]))
        Document(output)

    def test_legacy_table_cell_id_still_targets_the_first_paragraph(self):
        structure = DocumentStructureExtractor(str(self.source), self.logger).extract()
        output = self.root / "legacy.docx"

        suggestion = SuggestedChange(
            element_type="table_cell",
            element_id="table_1_cell_1_1",
            footnote_id=None,
            old="Första",
            new="Inledande",
            motivation="Bakåtkompatibilitet",
            match_status="exact",
        )
        plans = ChangePlanner(
            structure,
            TokenDiffEngine(logger=self.logger),
            self.logger,
        ).build_plans([suggestion])

        with DocxPackage(str(self.source), self.logger) as package:
            results = SimpleMarkupRenderer(package, self.logger).apply_plans(plans)
            package.save(str(output))

        self.assertTrue(results[0].applied, results[0].message)
        self.assertIn("Inledande", self._all_text(self._table_paragraphs(output)[0]))

    def _render_second_paragraph(self, output: Path, renderer_type):
        structure = DocumentStructureExtractor(str(self.source), self.logger).extract()
        suggestion = SuggestedChange(
            element_type="table_cell",
            element_id="table_1_cell_1_1_p2",
            footnote_id=None,
            old="måltexten",
            new="den tydliga texten",
            motivation="Regressionstest",
            match_status="exact",
        )
        plans = ChangePlanner(
            structure,
            TokenDiffEngine(logger=self.logger),
            self.logger,
        ).build_plans([suggestion])
        self.assertEqual(len(plans), 1)

        with DocxPackage(str(self.source), self.logger) as package:
            results = renderer_type(package, self.logger).apply_plans(plans)
            package.save(str(output))
        return results

    @staticmethod
    def _table_paragraphs(path: Path):
        with ZipFile(path) as archive:
            tree = etree.fromstring(archive.read("word/document.xml"))
        return tree.xpath("//w:body/w:tbl[1]/w:tr[1]/w:tc[1]/w:p", namespaces=NS)

    @staticmethod
    def _all_text(element):
        return "".join(element.xpath(".//w:t/text() | .//w:delText/text()", namespaces=NS))


if __name__ == "__main__":
    unittest.main()
