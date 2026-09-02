import logging
import os
import tempfile
import unittest
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile, ZipInfo

from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from lxml import etree

from app.src.JBGChangePlanner import ChangePlanner
from app.src.JBGCommentsRenderer import CommentsRenderer
from app.src.JBGDocumentStructureExtractor import DocumentStructureExtractor
from app.src.JBGDocxPackage import W_NS, DocxPackage
from app.src.JBGLangImprovSuggestorAI import SuggestedChange
from app.src.JBGSimpleMarkupRenderer import SimpleMarkupRenderer
from app.src.JBGTokenDiffEngine import TokenDiffEngine
from app.src.JBGTrackedChangesRenderer import TrackedChangesRenderer


NS = {"w": W_NS}


class HeaderFooterPipelineTests(unittest.TestCase):
    def setUp(self):
        self.temp_dir = tempfile.TemporaryDirectory()
        self.root = Path(self.temp_dir.name)
        self.logger = logging.getLogger(f"header-footer-test-{id(self)}")
        self.logger.handlers.clear()
        self.logger.addHandler(logging.NullHandler())
        self.logger.setLevel(logging.DEBUG)
        self.source = self.root / "source.docx"
        self._create_story_fixture(self.source)

    def tearDown(self):
        self.temp_dir.cleanup()

    def test_extracts_real_parts_once_and_resolves_inherited_sections(self):
        self._rename_story_part(
            self.source,
            old_part="word/header1.xml",
            new_part="word/storyParts/customHeader.xml",
        )

        structure = DocumentStructureExtractor(str(self.source), self.logger).extract()
        story_elements = [
            element
            for element in structure["elements"]
            if element["type"] in {"header", "footer"}
        ]

        self.assertTrue(story_elements)
        self.assertEqual(
            len({element["element_id"] for element in story_elements}),
            len(story_elements),
        )
        self.assertTrue(any(
            element["part_name"] == "word/storyParts/customHeader.xml"
            for element in story_elements
        ))

        default_header = self._element_containing(
            structure,
            element_type="header",
            text="Header alpha",
        )
        self.assertEqual(default_header["part_name"], "word/storyParts/customHeader.xml")
        self.assertEqual(default_header["section_indices"], [1, 2])
        self.assertEqual(default_header["story_variant"], "default")
        self.assertTrue(default_header["may_contain_special_runs"])

        first_header = self._element_containing(
            structure,
            element_type="header",
            text="First alpha",
        )
        even_header = self._element_containing(
            structure,
            element_type="header",
            text="Even alpha",
        )
        table_header = self._element_containing(
            structure,
            element_type="header",
            text="Header table",
        )
        self.assertEqual(first_header["section_indices"], [1, 2])
        self.assertEqual(even_header["section_indices"], [1, 2])
        self.assertIn("(.//w:p)", table_header["container_path"])

    def test_simple_markup_renders_header_and_footer_and_preserves_fields(self):
        self._rename_story_part(
            self.source,
            old_part="word/header1.xml",
            new_part="word/storyParts/customHeader.xml",
        )
        output = self.root / "simple.docx"
        before_field_count = self._count_nodes(
            self.source,
            "word/storyParts/customHeader.xml",
            ".//w:fldChar",
        )

        results = self._render(
            self.source,
            output,
            renderer_type=SimpleMarkupRenderer,
        )

        self.assertTrue(results)
        self.assertTrue(all(result.applied for result in results), [
            result.message for result in results
        ])
        self.assertEqual(
            self._count_nodes(output, "word/storyParts/customHeader.xml", ".//w:fldChar"),
            before_field_count,
        )

        header_tree = self._read_part_tree(output, "word/storyParts/customHeader.xml")
        footer_part = self._find_part_with_text(output, "word/footer", "Footer beta")
        footer_tree = self._read_part_tree(output, footer_part)
        self.assertTrue(header_tree.xpath(".//w:strike", namespaces=NS))
        self.assertTrue(footer_tree.xpath(".//w:strike", namespaces=NS))
        self.assertIn("Header gamma", self._all_text(header_tree))
        self.assertIn("Footer delta", self._all_text(footer_tree))

    def test_tracked_changes_render_header_and_footer_and_preserve_fields(self):
        output = self.root / "tracked.docx"
        header_part = self._find_part_with_text(self.source, "word/header", "Header alpha")
        before_field_count = self._count_nodes(
            self.source,
            header_part,
            ".//w:fldChar",
        )

        results = self._render(
            self.source,
            output,
            renderer_type=TrackedChangesRenderer,
        )

        self.assertTrue(all(result.applied for result in results), [
            result.message for result in results
        ])
        self.assertEqual(
            self._count_nodes(output, header_part, ".//w:fldChar"),
            before_field_count,
        )

        header_tree = self._read_part_tree(output, header_part)
        footer_part = self._find_part_with_text(output, "word/footer", "Footer beta")
        footer_tree = self._read_part_tree(output, footer_part)
        self.assertTrue(header_tree.xpath(".//w:del", namespaces=NS))
        self.assertTrue(header_tree.xpath(".//w:ins", namespaces=NS))
        self.assertTrue(footer_tree.xpath(".//w:del", namespaces=NS))
        self.assertTrue(footer_tree.xpath(".//w:ins", namespaces=NS))
        self.assertIn("Header gamma", self._all_text(header_tree))
        self.assertIn("Footer delta", self._all_text(footer_tree))

        settings_tree = self._read_part_tree(output, "word/settings.xml")
        self.assertTrue(settings_tree.xpath(".//w:trackRevisions", namespaces=NS))
        # A basic package round-trip catches broken relationships/content types.
        Document(output)

    def test_comments_anchor_to_tracked_changes_in_header_and_footer(self):
        output = self.root / "tracked-comments.docx"
        structure = DocumentStructureExtractor(str(self.source), self.logger).extract()
        suggestions = []
        for element_type, old, new in (
            ("header", "Header alpha", "Header gamma"),
            ("footer", "Footer beta", "Footer delta"),
        ):
            element = self._element_containing(structure, element_type, old)
            suggestions.append(SuggestedChange(
                element_type=element_type,
                element_id=element["element_id"],
                footnote_id=None,
                old=old,
                new=new,
                motivation=f"Why {element_type} changed",
                match_status="exact",
            ))

        plans = ChangePlanner(
            structure=structure,
            diff_engine=TokenDiffEngine(logger=self.logger),
            logger=self.logger,
        ).build_plans(suggestions)

        with DocxPackage(str(self.source), self.logger) as package:
            render_results = TrackedChangesRenderer(
                package,
                self.logger,
            ).apply_plans(plans)
            comment_results = CommentsRenderer(
                package,
                self.logger,
            ).apply_comments_for_results(render_results)
            package.save(str(output))

        self.assertTrue(all(result.applied for result in comment_results), [
            result.message for result in comment_results
        ])
        comments_tree = self._read_part_tree(output, "word/comments.xml")
        self.assertEqual(len(comments_tree.xpath(".//w:comment", namespaces=NS)), 2)

        for prefix, changed_text in (
            ("word/header", "Header gamma"),
            ("word/footer", "Footer delta"),
        ):
            part_name = self._find_part_with_text(output, prefix, changed_text)
            story_tree = self._read_part_tree(output, part_name)
            self.assertTrue(story_tree.xpath(".//w:commentRangeStart", namespaces=NS))
            self.assertTrue(story_tree.xpath(".//w:commentRangeEnd", namespaces=NS))
            self.assertTrue(story_tree.xpath(".//w:commentReference", namespaces=NS))

        Document(output)

    def _render(self, source: Path, output: Path, renderer_type):
        structure = DocumentStructureExtractor(str(source), self.logger).extract()
        changes = []
        for element_type, old, new in (
            ("header", "Header alpha", "Header gamma"),
            ("footer", "Footer beta", "Footer delta"),
        ):
            element = self._element_containing(structure, element_type, old)
            changes.append(SuggestedChange(
                element_type=element_type,
                element_id=element["element_id"],
                footnote_id=None,
                old=old,
                new=new,
                motivation="Regression test",
                match_status="exact",
            ))

        plans = ChangePlanner(
            structure=structure,
            diff_engine=TokenDiffEngine(logger=self.logger),
            logger=self.logger,
        ).build_plans(changes)
        self.assertEqual(len(plans), 2)

        with DocxPackage(str(source), self.logger) as package:
            results = renderer_type(package, self.logger).apply_plans(plans)
            package.save(str(output))
        return results

    @staticmethod
    def _create_story_fixture(path: Path):
        document = Document()
        document.add_paragraph("Body text")
        section = document.sections[0]

        header = section.header
        paragraph = header.paragraphs[0]
        paragraph.add_run("Header alpha ").bold = True
        HeaderFooterPipelineTests._append_page_field(paragraph)
        paragraph.add_run(" omega")
        table = header.add_table(rows=1, cols=1, width=Inches(2))
        table.cell(0, 0).text = "Header table"

        section.different_first_page_header_footer = True
        section.first_page_header.paragraphs[0].text = "First alpha"
        document.settings.odd_and_even_pages_header_footer = True
        section.even_page_header.paragraphs[0].text = "Even alpha"

        section.footer.paragraphs[0].text = "Footer beta"

        second = document.add_section(WD_SECTION.NEW_PAGE)
        second.header.is_linked_to_previous = True
        second.first_page_header.is_linked_to_previous = True
        second.even_page_header.is_linked_to_previous = True
        second.footer.is_linked_to_previous = True
        document.save(path)

    @staticmethod
    def _append_page_field(paragraph):
        begin_run = paragraph.add_run()
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instruction = OxmlElement("w:instrText")
        instruction.set(qn("xml:space"), "preserve")
        instruction.text = " PAGE "
        begin_run._r.append(begin)
        begin_run._r.append(instruction)

        result_run = paragraph.add_run("1")
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        result_run._r.append(end)

    @staticmethod
    def _element_containing(structure, element_type, text):
        matches = [
            element
            for element in structure["elements"]
            if element["type"] == element_type and text in element["text"]
        ]
        if len(matches) != 1:
            raise AssertionError(
                f"Expected one {element_type} containing {text!r}, found {len(matches)}"
            )
        return matches[0]

    @staticmethod
    def _rename_story_part(docx_path: Path, old_part: str, new_part: str):
        rewritten = docx_path.with_suffix(".rewritten.docx")
        old_target = os.path.basename(old_part)
        new_target = os.path.relpath(new_part, "word").replace(os.sep, "/")

        with ZipFile(docx_path, "r") as source, ZipFile(
            rewritten,
            "w",
            compression=ZIP_DEFLATED,
        ) as destination:
            for item in source.infolist():
                data = source.read(item.filename)
                name = new_part if item.filename == old_part else item.filename

                if item.filename == "word/_rels/document.xml.rels":
                    tree = etree.fromstring(data)
                    for relationship in tree:
                        if relationship.get("Target") == old_target:
                            relationship.set("Target", new_target)
                    data = etree.tostring(
                        tree,
                        xml_declaration=True,
                        encoding="UTF-8",
                    )
                elif item.filename == "[Content_Types].xml":
                    tree = etree.fromstring(data)
                    for override in tree:
                        if override.get("PartName") == f"/{old_part}":
                            override.set("PartName", f"/{new_part}")
                    data = etree.tostring(
                        tree,
                        xml_declaration=True,
                        encoding="UTF-8",
                    )

                replacement = ZipInfo(name, date_time=item.date_time)
                replacement.compress_type = ZIP_DEFLATED
                replacement.external_attr = item.external_attr
                destination.writestr(replacement, data)

        os.replace(rewritten, docx_path)

    @staticmethod
    def _read_part_tree(docx_path: Path, part_name: str):
        with ZipFile(docx_path) as archive:
            return etree.ElementTree(etree.fromstring(archive.read(part_name)))

    @classmethod
    def _count_nodes(cls, docx_path: Path, part_name: str, xpath: str):
        tree = cls._read_part_tree(docx_path, part_name)
        return len(tree.xpath(xpath, namespaces=NS))

    @classmethod
    def _find_part_with_text(cls, docx_path: Path, prefix: str, text: str):
        with ZipFile(docx_path) as archive:
            for name in archive.namelist():
                if name.startswith(prefix) and name.endswith(".xml"):
                    tree = etree.ElementTree(etree.fromstring(archive.read(name)))
                    if text in cls._all_text(tree):
                        return name
        raise AssertionError(f"Could not find {prefix} part containing {text!r}")

    @staticmethod
    def _all_text(tree):
        return "".join(tree.xpath(".//w:t/text() | .//w:delText/text()", namespaces=NS))


if __name__ == "__main__":
    unittest.main()
